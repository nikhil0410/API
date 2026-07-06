"""Generate a starter submission report (.docx) for the assignment.

Builds a Word document with the required screenshots-and-explanation structure:
title page, group/contribution tables, and a section per sub-objective with the
generated charts embedded and labeled placeholders for the screenshots that must
be captured live (Prefect dashboard, logs, API output).

Usage:
    python generate_report.py
Output:
    outputs/reports/Group<NN>_Submission_Report.docx
"""
from __future__ import annotations

import json

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import config

PLOTS = config.PLOTS_DIR
REPORTS = config.REPORTS_DIR


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, caption, width=5.8):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
    else:
        placeholder(doc, f"[MISSING IMAGE] {caption} (run the pipeline first)")


def placeholder(doc, text):
    """A visible grey box for a screenshot the student must insert."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("  " + text + "  ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x00)
    run.font.size = Pt(10)


def screenshot_slot(doc, what):
    placeholder(doc, f"[INSERT SCREENSHOT: {what}]")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("(replace this line with your captured screenshot)")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
def build() -> str:
    doc = Document()

    # ---- Title page ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AIMLCZG549 — API-driven Cloud Native Solutions")
    r.bold = True
    r.font.size = Pt(18)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("Assignment I — Cloud-Native ML Application\n"
                   "Predicting Student Dropout & Academic Success")
    rs.bold = True
    rs.font.size = Pt(14)

    para(doc, "Dataset: UCI Predict Students' Dropout and Academic Success "
              "(4,424 records, 36 features, binary Dropout-vs-not target).",
         italic=True, size=10).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Group details ----
    heading(doc, "Group Details", level=2)
    para(doc, "Group No: ______", bold=True)
    tbl = doc.add_table(rows=7, cols=5)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Sl. No", "BITS ID", "Name",
                           "Contribution (Qualitative)",
                           "% Contribution"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for i in range(1, 7):
        tbl.rows[i].cells[0].text = str(i)

    doc.add_page_break()

    # ---- 1. Business Understanding ----
    heading(doc, "1. Business Understanding", level=1)
    para(doc, "Higher-education institutions lose revenue and reputation when "
              "students drop out. Using academic, demographic and socio-economic "
              "attributes recorded at enrollment and after the first two "
              "semesters, this application predicts whether a student will "
              "Dropout, remain Enrolled, or Graduate — enabling early "
              "intervention by academic advisors.")
    bullet(doc, "Problem type: multi-class classification (3 classes).")
    bullet(doc, "Source: UCI ML Repository, dataset #697.")
    bullet(doc, "Records: 4,424  |  Features: 36  |  Target: Dropout / "
                "Enrolled / Graduate.")

    # ---- 2. Sub-Objective 1: Data Pipeline ----
    heading(doc, "2. Sub-Objective 1 — Data Pipeline (8 marks)", level=1)

    heading(doc, "2.1 Data Ingestion", level=2)
    para(doc, "The semicolon-delimited CSV downloaded from the UCI repository is "
              "loaded with pandas (data_pipeline.ingest_data). Column headers are "
              "stripped of trailing whitespace. The pipeline confirms 4,424 rows "
              "× 37 columns on ingestion.")
    screenshot_slot(doc, "terminal output of `python data_pipeline.py` showing "
                         "the 'Ingested 4424 rows' log line")

    heading(doc, "2.2 Data Pre-processing", level=2)
    para(doc, "Implemented in data_pipeline.preprocess_data:")
    bullet(doc, "Summary statistics for all columns saved to "
                "reports/summary_statistics.csv.")
    bullet(doc, "Missing-value check across all columns (dataset has 0 missing).")
    bullet(doc, "Median imputation for any numeric column with missing values.")
    bullet(doc, "Data types displayed (29 int64, 7 float64, 1 string target).")
    bullet(doc, "Min-Max normalization of all 36 numeric feature columns to [0, 1].")

    heading(doc, "2.3 Exploratory Data Analysis (EDA)", level=2)
    para(doc, "Implemented in data_pipeline.run_eda: correlation coefficients, "
              "label encoding, quartile binning of age, RandomForest feature "
              "importance, and univariate/bivariate visualizations.")
    add_image(doc, PLOTS / "target_distribution.png",
              "Figure 1: Univariate — target class distribution.")
    add_image(doc, PLOTS / "age_distribution.png",
              "Figure 2: Univariate — age at enrollment (normalized).")
    add_image(doc, PLOTS / "correlation_heatmap.png",
              "Figure 3: Correlation heatmap of the most relevant features.")
    add_image(doc, PLOTS / "feature_importance.png",
              "Figure 4: RandomForest feature importance (top 12).")
    add_image(doc, PLOTS / "bivariate_top_feature.png",
              "Figure 5: Bivariate — most important feature vs target.")

    # Top correlations from EDA report, if available
    eda_path = REPORTS / "eda_report.json"
    if eda_path.exists():
        eda = json.loads(eda_path.read_text())
        top = eda.get("top_correlated_features", {})
        if top:
            para(doc, "Top features correlated with the target:", bold=True)
            for k, v in list(top.items())[:5]:
                bullet(doc, f"{k}: |r| = {v}")

    heading(doc, "2.4 DataOps — Automation, Scheduling & Cloud Dashboard", level=2)
    para(doc, "The data pipeline is orchestrated with Prefect (flows.py). Tasks "
              "ingest → preprocess → EDA are wrapped as Prefect tasks whose logs "
              "are streamed to the Prefect dashboard. deploy.py registers a "
              "deployment with a 120-second (2-minute) interval schedule, so the "
              "workflow runs automatically every 2 minutes.")
    bullet(doc, "Start dashboard: `prefect server start` (http://127.0.0.1:4200).")
    bullet(doc, "Serve 2-minute schedule: `python deploy.py`.")
    screenshot_slot(doc, "Prefect dashboard — Flow Runs page showing successful "
                         "runs triggered every 2 minutes")
    screenshot_slot(doc, "Prefect dashboard — Deployments page showing the "
                         "'data-pipeline-every-2-min' deployment with its interval schedule")
    screenshot_slot(doc, "Prefect dashboard — Logs of a flow run showing all "
                         "activity (ingest / preprocess / EDA log lines)")

    doc.add_page_break()

    # ---- 3. Sub-Objective 2: ML Pipeline ----
    heading(doc, "3. Sub-Objective 2 — Machine Learning Pipeline (5 marks)", level=1)

    heading(doc, "3.1 Model Preparation", level=2)
    para(doc, "Two algorithms suited to tabular classification were "
              "selected: Random Forest and XGBoost (ml_pipeline.py). The feature "
              "matrix uses all 36 features (continuous imputed, low-cardinality "
              "dummy-encoded, high-cardinality Weight-of-Evidence encoded); the "
              "target is binary (Dropout=1 vs Not-Dropout=0).")

    heading(doc, "3.2 Model Training (70/30 split)", level=2)
    para(doc, "The dataset is split into 70% training and 30% testing with "
              "stratification (3,096 train / 1,328 test rows). Both models are "
              "tuned with RandomizedSearchCV; XGBoost uses a deeper, dedicated "
              "search (8 hyper-parameters, 30 candidates).")
    screenshot_slot(doc, "terminal output of `python ml_pipeline.py` showing the "
                         "70/30 split and 'Trained RandomForest / XGBoost' log lines")

    heading(doc, "3.3 Model Evaluation & 3.4 MLOps Metrics", level=2)
    para(doc, "Each model is evaluated with four metrics (accuracy, precision, "
              "recall, F1 for the Dropout class). MLOps logging writes each "
              "metric individually, appends a timestamped record to "
              "reports/model_metrics.jsonl, and tracks every run (params, "
              "metrics and artifacts) in MLflow under outputs/mlruns.")

    # Metrics table (read from latest run)
    metrics_path = REPORTS / "model_metrics_latest.json"
    if metrics_path.exists():
        data = json.loads(metrics_path.read_text())["metrics"]
        mtbl = doc.add_table(rows=1, cols=5)
        mtbl.style = "Light Grid Accent 1"
        mtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = mtbl.rows[0].cells
        for i, h in enumerate(["Model", "Accuracy", "Precision",
                               "Recall", "F1-score"]):
            hdr[i].paragraphs[0].add_run(h).bold = True
        for model, m in data.items():
            row = mtbl.add_row().cells
            row[0].text = model
            row[1].text = str(m["accuracy"])
            row[2].text = str(m["precision"])
            row[3].text = str(m["recall"])
            row[4].text = str(m["f1_score"])
        para(doc, "Table 1: Evaluation metrics on the held-out test set.",
             italic=True, size=9)

    add_image(doc, PLOTS / "confusion_matrix_RandomForest.png",
              "Figure 6: Confusion matrix — Random Forest.", width=4.5)
    add_image(doc, PLOTS / "confusion_matrix_XGBoost.png",
              "Figure 7: Confusion matrix — XGBoost.", width=4.5)
    screenshot_slot(doc, "terminal output showing the [MLOps] metric log lines "
                         "for both models")

    doc.add_page_break()

    # ---- 4. Sub-Objective 3: API Access ----
    heading(doc, "4. Sub-Objective 3 — API Access (2 marks)", level=1)
    para(doc, "api_access.py uses Prefect's built-in REST API (via the Python "
              "client) to retrieve and display key application details. At least "
              "four details are presented:")
    bullet(doc, "API health and API version.")
    bullet(doc, "Registered flows (e.g., data-pipeline, student-success-pipeline).")
    bullet(doc, "Deployments with their schedules and tags.")
    bullet(doc, "Recent flow runs and their states (Completed/Failed).")
    bullet(doc, "Work pools (cloud-native execution infrastructure).")
    para(doc, "Run: `python api_access.py` (with the Prefect server running).")
    screenshot_slot(doc, "terminal output of `python api_access.py` showing the "
                         "6 application details retrieved via the built-in API")

    # ---- 5. Conclusion ----
    heading(doc, "5. Conclusion", level=1)
    para(doc, "The application delivers an end-to-end cloud-native solution: an "
              "automated, scheduled and logged data pipeline (DataOps); a "
              "two-algorithm ML pipeline with monitored metrics (MLOps); and "
              "programmatic access to application details through built-in APIs. "
              "Random Forest and XGBoost are compared on the binary dropout "
              "prediction task (Dropout vs Not-Dropout).")

    out = REPORTS / "Group00_Submission_Report.docx"
    doc.save(str(out))
    return str(out)


if __name__ == "__main__":
    path = build()
    print(f"Report generated: {path}")
